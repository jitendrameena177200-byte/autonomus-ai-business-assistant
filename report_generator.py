"""
report_generator.py
--------------------
Generates polished, real-world business reports (.docx) from either:
  1. Summaries produced by summarizer.py, or
  2. Freeform LLM-generated content given a report topic/prompt.

The output is a formatted Word document with a title page, section
headings, and a timestamp - ready to send to stakeholders.
"""

from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI

from .config import settings

REPORT_PROMPT = ChatPromptTemplate.from_template(
    "You are a business analyst at {company_name}. Write a professional report "
    "section titled '{section_title}' based on the following notes. Use clear "
    "prose, be specific, and avoid generic filler. Notes:\n\n{notes}"
)


class ReportGenerator:
    def __init__(self) -> None:
        settings.validate()
        self.llm = ChatOpenAI(
            model=settings.chat_model, api_key=settings.openai_api_key, temperature=0.4
        )
        self.chain = REPORT_PROMPT | self.llm | StrOutputParser()

    def _expand_section(self, section_title: str, notes: str) -> str:
        return self.chain.invoke(
            {
                "company_name": settings.company_name,
                "section_title": section_title,
                "notes": notes,
            }
        )

    def generate_report(
        self,
        title: str,
        sections: Dict[str, str],
        expand_with_llm: bool = True,
        output_filename: Optional[str] = None,
    ) -> Path:
        """
        Build a .docx report.

        sections: dict mapping section heading -> raw notes/summary text.
        expand_with_llm: if True, each section's notes are expanded into
            polished prose by the LLM before being written to the doc.
        """
        doc = DocxDocument()

        # --- Title page ---
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run(
            f"{settings.company_name}  |  Generated {datetime.now().strftime('%B %d, %Y')}"
        )
        sub_run.italic = True
        sub_run.font.size = Pt(12)

        doc.add_page_break()

        # --- Sections ---
        for heading, notes in sections.items():
            doc.add_heading(heading, level=1)
            content = self._expand_section(heading, notes) if expand_with_llm else notes
            doc.add_paragraph(content)
            doc.add_paragraph("")

        # --- Footer note ---
        footer = doc.add_paragraph()
        footer_run = footer.add_run(
            "This report was generated automatically by the Autonomous AI Business "
            "Assistant. Please verify figures against source systems before external distribution."
        )
        footer_run.italic = True
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        filename = output_filename or f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = settings.output_dir / filename
        doc.save(str(output_path))
        return output_path
